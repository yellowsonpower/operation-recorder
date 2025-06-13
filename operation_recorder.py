import time
import pyautogui
import keyboard
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import threading
from pynput import mouse


class OperationRecorder:
    def __init__(self):
        self.screenshots = []
        self.recording = False
        self.output_folder = "operation_screenshots"
        self.docx_path = "operation_document.docx"

        # 创建输出文件夹
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)

    def record_mouse_click(self, x, y, button, pressed):
        """记录鼠标点击事件"""
        if not self.recording or not pressed:
            return

        # 截图
        screenshot = pyautogui.screenshot()

        # 在截图上标记鼠标点击位置 - 使用黄色实心圆
        draw = ImageDraw.Draw(screenshot)
        # 绘制黄色实心圆作为标记
        draw.ellipse((x - 15, y - 15, x + 15, y + 15), fill="yellow")
        # 添加红色边框使标记更明显
        draw.ellipse((x - 15, y - 15, x + 15, y + 15), outline="red", width=2)
        # 保留位置文本信息
        draw.text((x + 20, y - 10), f"Click at ({x}, {y})", fill="red")

        # 保存截图
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        img_path = os.path.join(self.output_folder, f"screenshot_{timestamp}.png")
        screenshot.save(img_path)

        # 添加到截图列表
        self.screenshots.append((img_path, f"鼠标点击: ({x}, {y})"))

    def record_keyboard_event(self, e):
        """记录键盘事件"""
        if not self.recording:
            return

        # 获取当前鼠标位置
        x, y = pyautogui.position()

        # 截图
        screenshot = pyautogui.screenshot()

        # 在截图上标记键盘事件
        draw = ImageDraw.Draw(screenshot)
        draw.rectangle((x - 10, y - 10, x + 100, y + 20), outline="blue", width=2)
        draw.text((x + 10, y - 5), f"Key: {e.name}", fill="blue")

        # 保存截图
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        img_path = os.path.join(self.output_folder, f"screenshot_{timestamp}.png")
        screenshot.save(img_path)

        # 添加到截图列表
        self.screenshots.append((img_path, f"键盘按键: {e.name}"))

    def start_recording(self):
        """开始记录"""
        self.recording = True
        print("开始记录操作，请执行你的操作...")
        print("按 Ctrl+Shift+S 停止记录并生成Word文档")

        # 设置键盘监听
        keyboard.on_press(self.record_keyboard_event)

        # 设置鼠标监听
        self.mouse_listener = mouse.Listener(on_click=self.record_mouse_click)
        self.mouse_listener.start()

        # 等待停止信号
        keyboard.wait('ctrl+shift+s')
        self.stop_recording()

    def stop_recording(self):
        """停止记录"""
        self.recording = False
        self.mouse_listener.stop()
        print("停止记录，正在生成Word文档...")
        self.generate_docx()

    def generate_docx(self):
        """生成Word文档"""
        if not self.screenshots:
            print("没有记录到任何操作")
            return

        # 创建Word文档
        doc = Document()

        # 设置文档默认字体为微软雅黑
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 设置标题字体大小
        doc.styles['Heading 1'].font.size = Pt(16)
        doc.styles['Heading 1'].font.name = '微软雅黑'
        doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加标题
        title = doc.add_heading('操作记录文档', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加日期时间
        current_time = time.strftime("%Y-%m-%d %H:%M:%S")
        date_paragraph = doc.add_paragraph(f'记录时间: {current_time}')
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')  # 添加空行

        for img_path, description in self.screenshots:
            try:
                # 添加操作描述
                p = doc.add_paragraph()
                p.add_run(description).bold = True

                # 添加截图
                doc.add_picture(img_path, width=Inches(6))

                # 添加分隔线
                doc.add_paragraph('-' * 50)

            except Exception as e:
                print(f"处理图片 {img_path} 时出错: {e}")

        # 保存Word文档
        doc.save(self.docx_path)
        print(f"Word文档已生成: {os.path.abspath(self.docx_path)}")


if __name__ == "__main__":
    recorder = OperationRecorder()
    recorder.start_recording()
